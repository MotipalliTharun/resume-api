import os
import re
import subprocess
import tempfile
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# WeasyPrint removed per update

def set_run_font(run, font_name='Times New Roman', font_size=10.5, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    # Force font name for compatibility
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)

def create_resume_docx(data: dict, output_path: str):
    """
    Generates a structured ATS-optimized DOCX from the extracted data dictionary.
    """
    doc = Document()
    
    # 1. Page Setup (0.75 inch margins)
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    
    # 2. Styles (Times New Roman default)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10.5)
    
    # 3. Header (Name & Contact)
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Name
    if data.get("full_name"):
        run = p_name.add_run(data["full_name"].title())
        set_run_font(run, font_size=22, bold=True)
    
    # Contact Info line
    # content: "Location | Phone | Email | LinkedIn"
    if data.get("contact_info_block"):
        # Use raw extracted block if available (for updates/regenerations)
        # remove newlines to keep it one line if possible, or usually it's one line
        text = data["contact_info_block"].replace("\n", " | ").strip()
        
        # Deduplicate Name: Remove name if it appears in the contact block
        name = data.get("full_name", "").strip()
        if name:
             # Case-insensitive remove
             pattern = re.compile(re.escape(name), re.IGNORECASE)
             text = pattern.sub("", text).strip()
             
             # Cleanup leading separators
             text = text.lstrip(" |:-•").strip()

        p_contact = doc.add_paragraph(text)
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_contact.runs:
             set_run_font(run, font_size=10.5, bold=False)
        p_contact.paragraph_format.space_after = Pt(12)
        p_contact.paragraph_format.line_spacing = 1.15
    else:
        parts = []
        if data.get("location"): parts.append(data["location"])
        if data.get("phone"): parts.append(data["phone"])
        if data.get("email"): parts.append(data["email"])
        if data.get("linkedin"): parts.append(data["linkedin"])
        if data.get("portfolio"): parts.append(data["portfolio"])
        
        if parts:
            p_contact = doc.add_paragraph(" | ".join(parts))
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p_contact.runs:
                 set_run_font(run, font_size=10.5, bold=False)
            p_contact.paragraph_format.space_after = Pt(12)
            p_contact.paragraph_format.line_spacing = 1.15
    
    # 4. Helper to add Section
    
    def add_section(title, content):
        if not content: return
        
        # Section Heading
        h = doc.add_paragraph(title.upper())
        run = h.runs[0] if h.runs else h.add_run(title.upper())
        set_run_font(run, font_size=12, bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0) # Black
        
        # Border
        p_element = h._p
        pPr = p_element.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pbdr.append(bottom)
        pPr.append(pbdr)
        
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6) 
        h.paragraph_format.line_spacing = 1.0 
        
        # Content Parsing
        last_was_bullet = False
        
        lines = [l.strip() for l in content.splitlines() if l.strip()]
        
        for i, line in enumerate(lines):
            # 1. CLEANING: Remove typical bullet chars but PROTECT 'make bold' chars (**)
            # We want to remove '• ', '- ', '* ' but NOT '**'
            clean_line = line
            is_bullet = False
            
            # Check for standard bullets
            if line.startswith("•") or line.startswith("-") or (line.startswith("*") and not line.startswith("**")):
                # potential bullet
                clean_line = line.lstrip("•- ").strip()
                if clean_line.startswith("*") and not clean_line.startswith("**"):
                     clean_line = clean_line.lstrip("*").strip()
                is_bullet = True
            
            # 2. HEURISTIC: Promote "Bold Bullets" to Headers
            # If a line is a bullet, but it seems to be entirely bold (e.g. **Job Title**)
            # We treat it as a Header (Paragraph) instead of a List Item.
            # This fixes cases where AI adds bullets to headers.
            is_entirely_bold = clean_line.startswith("**") and clean_line.endswith("**") and len(clean_line) > 4
            
            if is_bullet and is_entirely_bold:
                is_bullet = False # Demote from bullet to header
            
            # 3. RENDER
            if is_bullet:
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                last_was_bullet = True
            else:
                p = doc.add_paragraph()
                
                # Intelligent Spacing for Headers
                if last_was_bullet and i > 0:
                     p.paragraph_format.space_before = Pt(12) 
                else:
                     p.paragraph_format.space_before = Pt(2)
                
                p.paragraph_format.keep_with_next = True
                last_was_bullet = False
            
            # 4. ROBUST REGEX BOLD PARSING
            tokens = re.split(r'(\*\*.*?\*\*)', clean_line)
            
            for token in tokens:
                if not token: continue
                
                if token.startswith('**') and token.endswith('**') and len(token) > 4:
                    text = token[2:-2]
                    r = p.add_run(text)
                    set_run_font(r, font_size=10.5, bold=True)
                else:
                    r = p.add_run(token)
                    set_run_font(r, font_size=10.5, bold=False)
            
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15

    # 5. Add Sections in Order
    # Map data keys to Display Titles
    
    # Summary
    add_section("PROFESSIONAL SUMMARY", data.get("summary"))
    
    # Skills
    add_section("CORE SKILLS", data.get("skills_block"))
    
    # Experience
    add_section("PROFESSIONAL EXPERIENCE", data.get("experience_block"))
    
    # Education
    add_section("EDUCATION", data.get("education_block"))
    
    # Projects
    add_section("PROJECTS", data.get("projects_block"))
    
    # Certifications
    add_section("CERTIFICATIONS", data.get("certifications_block"))
    
    # Achievements
    add_section("ACHIEVEMENTS", data.get("achievements_block"))
    
    # Volunteer
    add_section("VOLUNTEER", data.get("volunteer_block"))
    
    doc.save(output_path)


def convert_docx_to_pdf(docx_path: str, output_dir: str = None) -> str:
    """
    Converts a DOCX file to PDF using LibreOffice (soffice).
    Returns the path to the generated PDF.
    """
    input_path = Path(docx_path)
    if not input_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    if output_dir is None:
        output_dir = input_path.parent
    
    # soffice --headless --convert-to pdf <file> --outdir <dir>
    # Note: On Mac standard install, 'soffice' might not be in PATH. 
    # Logic to find or try common paths could be added here or relied on env.
    
    cmd = [
        "soffice", 
        "--headless", 
        "--convert-to", "pdf", 
        str(input_path), 
        "--outdir", str(output_dir)
    ]
    
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        pdf_path = Path(output_dir) / (input_path.stem + ".pdf")
        if pdf_path.exists():
            return str(pdf_path)
        else:
            print("PDF conversion failed: Output file not created.")
            return ""
    except Exception as e:
        print(f"PDF conversion error: {e}")
        # Identify if soffice is missing
        if "No such file or directory" in str(e):
            print("LibreOffice (soffice) not found. Please install it.")
        return ""

