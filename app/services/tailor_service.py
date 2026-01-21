import json
from pathlib import Path
from sqlalchemy.orm import Session
from models import TailorRun
from services.render_service import render_md, render_html
from utils import safe_filename

# Determine DATA_DIR relative to this file
# This file is in app/services/, so parent.parent is app/
DATA_DIR = Path(__file__).resolve().parent.parent / "_data"
DATA_DIR.mkdir(parents=True, exist_ok=True)

def extract_section(full_text: str, heading: str) -> str:
    lines = full_text.splitlines()
    # Normalize heading for match (strip markdown chars like *, #)
    h_norm = heading.strip().upper().replace("*", "").replace("#", "")
    
    start = None
    for i, ln in enumerate(lines):
        ln_norm = ln.strip().upper().replace("*", "").replace("#", "")
        if ln_norm == h_norm:
            start = i + 1
            break
            
    if start is None:
        return ""
        
    end = len(lines)
    for j in range(start, len(lines)):
        s = lines[j].strip()
        # End of section heuristic: another potential header
        if s and s.upper() == s and len(s.replace("*", "")) <= 30:
            # Check if this looks like another header
            s_norm = s.replace("*", "").replace("#", "").strip()
            # If it's a known header or looks like one, stop
            if s_norm in ["SUMMARY", "SKILLS", "WORK EXPERIENCE", "EDUCATION", "PROJECTS", "CERTIFICATIONS", "ACHIEVEMENTS", "VOLUNTEER", "PERSONAL DETAILS"]:
                end = j
                break
    return "\n".join(lines[start:end]).strip()

def process_and_save_run(
    db: Session,
    full_text: str,
    run_context: dict
) -> TailorRun:
    """
    Common logic to parse AI output, render templates, save to DB, and write files.
    """
    first_line = full_text.splitlines()[0] if full_text.splitlines() else ""
    
    # --- PROGRAMMATIC KEYWORD NORMALIZATION ---
    # Enforce JD terminology (ReactJS vs React, AWS vs aws)
    try:
        from services.keyword_normalization_service import KeywordStandardizer
        jd = run_context.get("jd_text", "")
        full_text = KeywordStandardizer.normalize_text(full_text, jd)
    except Exception as e:
        print(f"Normalization Error (non-fatal): {e}")
    # ------------------------------------------

    data = {
        "full_name": run_context.get("full_name") or "YOUR NAME",
        "location": run_context.get("location") or "",
        "phone": run_context.get("phone") or "",
        "email": run_context.get("email") or "",
        "linkedin": run_context.get("linkedin") or "",
        "portfolio": run_context.get("portfolio") or "",
        "summary": extract_section(full_text, "SUMMARY"),
        "skills_block": extract_section(full_text, "SKILLS"),
        "experience_block": extract_section(full_text, "WORK EXPERIENCE"),
        "projects_block": extract_section(full_text, "PROJECTS"), 
        "certifications_block": extract_section(full_text, "CERTIFICATIONS"),
        "achievements_block": extract_section(full_text, "ACHIEVEMENTS"), 
        "volunteer_block": extract_section(full_text, "VOLUNTEER"),
        "education_block": extract_section(full_text, "EDUCATION"),
    }

    tailored_md = render_md("templates/ats_resume.md", data)
    tailored_html = render_html("templates/ats_resume.html", data)

    run = TailorRun(
        job_title=run_context.get("job_title", ""),
        company=run_context.get("company", ""),
        full_name=run_context.get("full_name", ""),
        jd_text=run_context.get("jd_text", ""),
        resume_text=run_context.get("resume_text", ""),
        score_keyword=float(run_context.get("kw_score", 0)),
        score_semantic=float(run_context.get("sem_score", 0)),
        score_ats=float(run_context.get("ats_score", 0)),
        score_total=float(run_context.get("total", 0)),
        missing_keywords=",".join(run_context.get("missing", [])[:60]),
        tailored_text=full_text,
        tailored_md=tailored_md,
        tailored_html=tailored_html,
    )
    
    db.add(run)
    db.commit()
    db.refresh(run)

    # Save artifacts
    base = f"run_{run.id}_{safe_filename(run.job_title or 'role')}"
    (DATA_DIR / f"{base}.txt").write_text(full_text, encoding="utf-8")
    (DATA_DIR / f"{base}.md").write_text(tailored_md, encoding="utf-8")
    (DATA_DIR / f"{base}.html").write_text(tailored_html, encoding="utf-8")

    # Save DOCX artifact
    from services.export_service import create_resume_docx, convert_docx_to_pdf
    
    docx_path = DATA_DIR / f"{base}.docx"
    create_resume_docx(data, str(docx_path))
    
    # Generate PDF from DOCX
    pdf_path = convert_docx_to_pdf(str(docx_path))
    if pdf_path:
        print(f"PDF Generated: {pdf_path}")
    
    return run

def update_run_text(db: Session, run: TailorRun, new_text: str):
    """
    Updates the run's tailored text and re-renders MD/HTML artifacts.
    """
    # Extract data using existing sections or keep original contact info
    data = {
        "full_name": run.full_name,
        # Extract Contact Info directly from text to avoid losing it during updates
        "contact_info_block": extract_section(new_text, "PERSONAL DETAILS"),
        "location": "", 
        "phone": "",
        "email": "",
        "linkedin": "",
        "portfolio": "",
        "summary": extract_section(new_text, "SUMMARY"),
        "skills_block": extract_section(new_text, "SKILLS"),
        "experience_block": extract_section(new_text, "WORK EXPERIENCE"),
        "projects_block": extract_section(new_text, "PROJECTS"),
        "certifications_block": extract_section(new_text, "CERTIFICATIONS"),
        "achievements_block": extract_section(new_text, "ACHIEVEMENTS"),
        "education_block": extract_section(new_text, "EDUCATION"),
        "volunteer_block": extract_section(new_text, "VOLUNTEER"),
    }

    run.tailored_text = new_text
    run.tailored_md = render_md("templates/ats_resume.md", data)
    run.tailored_html = render_html("templates/ats_resume.html", data)

    db.commit()
    db.refresh(run)

    # Overwrite artifacts
    base = f"run_{run.id}_{safe_filename(run.job_title or 'role')}"
    (DATA_DIR / f"{base}.txt").write_text(new_text, encoding="utf-8")
    (DATA_DIR / f"{base}.md").write_text(run.tailored_md, encoding="utf-8")
    (DATA_DIR / f"{base}.html").write_text(run.tailored_html, encoding="utf-8")
    
    # REGENERATE DOCX & PDF
    from services.export_service import create_resume_docx, convert_docx_to_pdf
    
    docx_path = DATA_DIR / f"{base}.docx"
    # Remove old files to ensure fresh generation (optional but safe)
    if docx_path.exists(): docx_path.unlink()
    
    create_resume_docx(data, str(docx_path))
    
    pdf_path = convert_docx_to_pdf(str(docx_path))
    if pdf_path:
        print(f"Updated PDF: {pdf_path}")
    
    return run


def process_and_save_cover_letter(
    db: Session,
    full_text: str,
    run_context: dict
) -> TailorRun:
    from models import CoverLetter
    import datetime
    
    # Prepare data for template
    from services.render_service import md_filter
    body_html = md_filter(full_text)
    
    data = {
        "full_name": run_context.get("full_name") or "Your Name",
        "phone": run_context.get("phone") or "",
        "email": run_context.get("email") or "",
        "linkedin": run_context.get("linkedin") or "",
        "location": run_context.get("location") or "",
        "portfolio": run_context.get("portfolio") or "",
        "current_date": datetime.date.today().strftime("%B %d, %Y"),
        "hiring_manager": "Hiring Manager",
        "company": run_context.get("company", ""),
        "company_address": "",
        "body_html": body_html
    }

    # Render Artifacts
    cover_letter_html = render_html("templates/cover_letter.html", data)

    # Save to DB
    cl = CoverLetter(
        job_title=run_context.get("job_title", ""),
        company=run_context.get("company", ""),
        full_name=run_context.get("full_name", ""),
        jd_text=run_context.get("jd_text", ""),
        resume_text=run_context.get("resume_text", ""),
        cover_letter_text=full_text,
        cover_letter_html=cover_letter_html
    )
    db.add(cl)
    db.commit()
    db.refresh(cl)

    # Save locally
    base = f"cl_{cl.id}_{safe_filename(cl.company or 'company')}"
    (DATA_DIR / f"{base}.txt").write_text(full_text, encoding="utf-8")
    (DATA_DIR / f"{base}.html").write_text(cover_letter_html, encoding="utf-8")

    # Save DOCX
    # Create simple but formatted DOCX
    from docx import Document
    from docx.shared import Pt, Inches
    from services.export_service import convert_docx_to_pdf, set_run_font

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Add text line by line
    for line in full_text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(line)
        set_run_font(run, font_name='Times New Roman', font_size=11)

    docx_path = DATA_DIR / f"{base}.docx"
    doc.save(str(docx_path))

    # Generate PDF
    convert_docx_to_pdf(str(docx_path))

    return cl